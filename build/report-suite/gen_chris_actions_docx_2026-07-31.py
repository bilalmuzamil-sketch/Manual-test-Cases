#!/usr/bin/env python3
"""Generate the PO-facing Word document "What We Need From Chris Ward".

Mirrors the house style of the existing leadership/PO-facing docs (Standing Rule 16):
  build/How-We-Ensure-Test-Case-Quality_Simple-Guide_2026-07-29.docx
  build/Test-Case-Creation-and-Refinement-Process_2026-07-29.docx
i.e. Calibri throughout, centred 22pt bold #1F4E79 title, 13pt #555555 subtitle,
10pt #777777 prepared line, 14pt bold #1F4E79 section headings, 11pt body,
bullet lists via "List Paragraph".

READER-FACING ONLY — the QA-internal appendix of the .md is deliberately NOT emitted
here, so the file can be sent to the PO as-is.

Run from anywhere:  python3 build/report-suite/gen_chris_actions_docx_2026-07-31.py
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

OUT = Path(__file__).resolve().parent / "What-We-Need-From-Chris-Ward-2026-07-31.docx"

NAVY = RGBColor(0x1F, 0x4E, 0x79)
GREY_MID = RGBColor(0x55, 0x55, 0x55)
GREY_LIGHT = RGBColor(0x77, 0x77, 0x77)


def para(doc, text="", size=11, bold=False, color=None, align=None,
         space_after=8, style=None, italic=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.paragraph_format.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def rich(doc, chunks, size=11, space_after=8, style=None):
    """chunks = [(text, bold), ...]"""
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    for text, bold in chunks:
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.bold = bold
    return p


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.bold = True
    r.font.color.rgb = NAVY
    return p


def item(doc, number, title, need, why, ifstays, bywhen):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"{number}. {title}")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.bold = True
    r.font.color.rgb = GREY_MID
    for label, body in (
        ("What we need from you: ", need),
        ("Why: ", why),
        ("What happens if it stays as it is: ", ifstays),
        ("By when: ", bywhen),
    ):
        rich(doc, [(label, True), (body, False)], size=11, space_after=3,
             style="List Paragraph")


# ---------------------------------------------------------------- content
DECISIONS = [
    (
        "Sales By Representative — do the downloaded files follow the screen?",
        "Tell us which is right — the location column appears in the four downloaded files "
        "whenever it is showing on screen, or the downloads always show one fixed set of columns.",
        "Your Sales By Representative description now says both. A newer line says the location "
        "column is in all four downloads; older lines list the download columns as a fixed set, "
        "in order, with no location column in it.",
        "Two readers of the same description get two different answers, and a tester at a company "
        "with more than one location can report a correct download as broken. We have followed the "
        "newer line for now.",
        "At your convenience — before anyone tests the downloads.",
    ),
    (
        "Where the location column sits in the two shorter downloads",
        "One line saying where the location column goes in the two shorter “Summary” downloads.",
        "Your instruction is that it sits in the same place it sits on screen. Those two files do "
        "not contain the column it sits next to on screen, so there is no matching place, and "
        "nothing anywhere says where it should go.",
        "Our check has to accept whatever position the build uses, so a wrong position would pass "
        "unnoticed.",
        "At your convenience.",
    ),
    (
        "One logo rule for all six printed downloads",
        "Pick the single logo rule every report follows.",
        "You told us on 29 July that every report now uses the same logo treatment, but the "
        "descriptions give three different rules: one says try the company’s own logo first, then "
        "the built-in ShopView one, and print none if neither exists; another says always print the "
        "built-in ShopView logo; a third does not mention a logo at all.",
        "Three of our checks follow three different rules, so at least two of them are wrong and "
        "nobody can tell which.",
        "At your convenience.",
    ),
    (
        "Does “normal reports access” mean one permission for all six?",
        "Confirm whether all six reports open with one single reports permission, or whether the "
        "existing per-area reports permissions (such as the inventory one) still apply.",
        "You have told us twice that these reports must not sit behind their own special "
        "permissions. Two of them — Parts Velocity and Inventory Value — are described as needing "
        "the existing inventory-reports permission, which is itself a normal one, so your "
        "instruction can honestly be read two ways.",
        "Whoever sets up a test user does not know which permission to give them, so “who can open "
        "this report” goes untested for two reports. Separately, the Sales By Customer description "
        "still says that report has its own special permission — that line needs correcting "
        "whichever way you answer.",
        "At your convenience.",
    ),
    (
        "Should the screens keep saying “VIN”, even for a generator?",
        "Confirm the screens keep the word “VIN” for every asset, or tell us the wording you want "
        "for assets that are not vehicles.",
        "You raised this yourself on 29 July — VIN means vehicle identification number, so for "
        "something like a generator the number in that field is really a serial number. Our checks "
        "currently expect the screens to keep saying “VIN”.",
        "Nothing breaks, but the reports will show “VIN” beside a generator’s serial number, and "
        "customers will ask about it.",
        "At your convenience.",
    ),
]

WRITEDOWN = [
    (
        "Work In Progress — put the VIN first (you believe this is already done)",
        "Edit the Work In Progress description so an asset is identified by its VIN first, then the "
        "unit number, then the plate.",
        "On 29 July you told us this is the standard for every report, and that you had already "
        "made this edit. The description still puts the unit number first in six places.",
        "On every screen that shows an asset, the description and our checks say the opposite of "
        "each other. This is the one to do first.",
        "4 August — the date we agreed.",
    ),
    (
        "The location dropdown disappears for a one-location person",
        "Correct the four descriptions — Sales By Representative, Technician Utilization, Inventory "
        "Value and Parts Velocity — which still say that someone with only one location still sees "
        "the location dropdown.",
        "You ruled on 31 July that it is hidden.",
        "Four descriptions actively state the opposite of your own ruling.",
        "4 August.",
    ),
    (
        "The full word “Representative” everywhere",
        "Replace “Sales Rep” with “Sales Representative” in the Sales By Representative "
        "description — the row on the customer record, the picker on the work order, and the "
        "assignments download (its menu entry, its file name and its column heading).",
        "You ruled on 31 July that “Rep” is out everywhere.",
        "The description still shows the short form in several places, so nobody reading it can "
        "tell what the screens are supposed to say.",
        "4 August.",
    ),
    (
        "The download size limit, and the message when it is hit",
        "Add the 10,000-row download limit to the Parts Velocity, Technician Utilization and Work "
        "In Progress descriptions, and correct the Sales By Customer wording so all six use the one "
        "message you chose: “This report is too large to export. Narrow the date range or filters, "
        "then try again.”",
        "You confirmed on 31 July that the limit applies to all six and that there is one message.",
        "Three descriptions say nothing about a limit at all, so anyone reading them would call our "
        "checks unfounded.",
        "4 August.",
    ),
    (
        "The Escape key on the “deactivate a sales representative” pop-up",
        "Correct the Sales By Representative description, which still says that pop-up closes when "
        "you press the Escape key.",
        "You ruled on 28 July that it should not — the app’s general house rule wins. Our check "
        "follows your ruling.",
        "The description asks for behaviour you have ruled out, and the developers are still "
        "carrying it as an unresolved question.",
        "At your convenience — this is the oldest item here, from 28 July.",
    ),
    (
        "Three small tidy-ups",
        "(a) In the Sales By Customer description, name the menu group these reports sit in and say "
        "the new links go below the existing ones. (b) Correct the Parts Velocity line calling it "
        "the “only” report in the Parts group — Inventory Value is there too. (c) Two descriptions, "
        "Sales By Representative and Parts Velocity, have a few characters showing as odd symbols.",
        "All three come from your own video and your own note.",
        "Small things, but each one is a reader asking us a question you have already answered.",
        "At your convenience.",
    ),
    (
        "The new “choose your columns” control on Technician Utilization has no ticket",
        "Have a ticket raised for the column-choosing control you asked for on Technician "
        "Utilization.",
        "You asked for it in your 29 July note and we have written the checks, but there is no "
        "piece of work in the tracker for it.",
        "Those two checks are the only ones in the whole set that cannot be tied back to a specific "
        "piece of work.",
        "At your convenience.",
    ),
]


def main():
    doc = Document()
    for s in doc.sections:
        s.left_margin = s.right_margin = Pt(72)
        s.top_margin = s.bottom_margin = Pt(54)

    para(doc, "What We Need From You — the six new reports", size=22, bold=True,
         color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    para(doc, "For Chris Ward, Product Owner", size=13, color=GREY_MID,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=3)
    para(doc, "Prepared 31 July 2026  ·  Five decisions and seven write-it-downs  ·  "
              "no technical knowledge needed", size=10, color=GREY_LIGHT,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    para(doc, "This is not a questionnaire. It is a short list of things we need you to do, and "
              "what each one unblocks.")
    para(doc, "Our testers check the reports against your written report descriptions. Where a "
              "description says one thing and you have told us another, a tester can report a "
              "perfectly good report as broken — or miss a real problem. That is the risk every "
              "item below removes.")
    rich(doc, [("Five things need your decision. Seven are answers you have already given us that "
                "have not yet reached the written descriptions.", True)])

    heading(doc, "Things only you can decide")
    for i, row in enumerate(DECISIONS, start=1):
        item(doc, i, *row)

    heading(doc, "Things that just need writing down")
    para(doc, "You have already answered all seven of these. Our checks follow your answers, not "
              "the older written text. What is missing is the edit to the descriptions — until it "
              "lands, the descriptions and our checks disagree, and anyone comparing the two will "
              "assume the mistake is ours.")
    for i, row in enumerate(WRITEDOWN, start=len(DECISIONS) + 1):
        item(doc, i, *row)

    para(doc, "")
    rich(doc, [("Thank you", True),
               (" — the 29 July round of description updates landed on time and cleared most of "
                "what we were waiting on, and your answers on 31 July settled five open points in "
                "a day. The list above is what is left.", False)])

    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
