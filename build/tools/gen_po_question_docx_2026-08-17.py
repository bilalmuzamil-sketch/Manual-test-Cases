#!/usr/bin/env python3
"""Generate Google-Docs-ready .docx versions of the two PO question sheets.
Reader-facing content ONLY. No case IDs, C-ids, SV- keys, spec anchors, VIU, or HTTP terms.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# CONTENT (reader-facing only; jargon deliberately excluded)
# ---------------------------------------------------------------------------

CHRIS = {
    "path": "build/report-suite/questions-2026-08-17/Report-Suite_Questions-for-Chris-Ward_2026-08-17.docx",
    "title": "Questions for Chris Ward",
    "subtitle": "Report Suite (the six reports) — Product Owner: Chris Ward",
    "intro": (
        "Two quick questions, both about the Work In Progress report. Each one is a plain "
        "A / B / C, and both are small tidy-ups where your written description says a thing "
        "two different ways and we would rather have your word than guess. Every question names "
        "the project and the report, because we know you look after more than one thing here. "
        "There are no bugs on this sheet — just two wording decisions. Thank you."
    ),
    "questions": [
        {
            "heading": "Report Suite — the Work In Progress report — the Estimates help text "
                       "(the little information icon next to the Estimates figure)",
            "now": [
                "On the Work In Progress report, the Estimates figure has a small information "
                "icon that shows a short plain explanation of what that number means.",
                "Your written description gives that explanation TWO different ways, in two "
                "different places of the same document — one short, and one longer that the "
                "recent design review locked in. Word for word:",
                "• Short version: “Quotes the customer has not approved yet — not counted "
                "in the totals.”",
                "• Longer version (locked in the design review): “The total value of all "
                "estimate lines that have not yet been approved, including lines awaiting "
                "authorization on open work orders.”",
                "We are using the longer one, because it is the most recent — it came from the "
                "design review. We just want to record which one you want, so the description "
                "says it only once.",
                "Why we are asking: it is a one-line tidy-up in your description either way, and "
                "we would rather have your word than leave the description saying two things.",
            ],
            "question": "Which explanation should the Estimates information icon show — and may "
                        "we drop the other so the description states it once?",
            "options": [
                "A) Keep the longer one from the design review (“The total value of all "
                "estimate lines that have not yet been approved, including lines awaiting "
                "authorization on open work orders”) and drop the short one. This is the "
                "one we have already built our check to, so if you choose A we just need your "
                "confirmation and the short line tidied out.",
                "B) Keep the short one (“Quotes the customer has not approved yet — not "
                "counted in the totals”) and drop the longer one. We change our check back "
                "to the short wording.",
                "C) Something else — please write the exact wording you want.",
            ],
        },
        {
            "heading": "Report Suite — the Work In Progress report — which tab a job appears in "
                       "(the tabs across the top of the report)",
            "now": [
                "The Work In Progress report is split into tabs across the top — Estimate, In "
                "Progress, Review, Complete, and so on.",
                "Your written description says two different things about which tab a job "
                "belongs in, in the same document:",
                "• One part says a job appears in exactly ONE tab, chosen by the job's overall "
                "status.",
                "• Another part, added more recently, says the tabs go by the state of each "
                "individual LINE on the job — so a job with lines in more than one state would "
                "show up in more than one tab.",
                "We have not picked a side — our tests follow the wording each was written "
                "against. This is the one thing on the report we cannot settle ourselves.",
                "Why we are asking: the two behave completely differently for a job that has "
                "work in more than one state, and we do not want to guess which one is right.",
            ],
            "question": "When a single job has work in more than one state, should it appear in "
                        "just one tab, or in every tab that matches?",
            "options": [
                "A) In just ONE tab, chosen by the job's overall status — a job is only ever in "
                "one place on the report.",
                "B) In EVERY tab that matches — a job with some lines estimated and some in "
                "progress shows in both the Estimate tab and the In Progress tab.",
                "C) Something else — please describe it.",
            ],
        },
    ],
}

# ---------------------------------------------------------------------------
# The five NEW follow-up questions (Aug-5 design review; Schedule spec v30 silent).
# Shared between the full Branko sheet (as questions 4-8) and the follow-up-only doc.
# ---------------------------------------------------------------------------
FOLLOWUP_5 = [
    {
        "heading": "SCHEDULE — the scheduling pop-up — a personal 'always schedule the whole work "
                   "order' preference",
        "now": [
            "When you drag a work order onto the schedule, it asks whether to schedule the WHOLE "
            "work order or to pick individual lines. Right now it asks this every time.",
            "The 5 August design review suggested letting each person set a preference — 'always "
            "schedule the whole work order' — so the system remembers their choice and stops "
            "asking every time.",
            "Your newer written description does not mention this preference at all, and the review "
            "itself listed it as an open question to decide before launch. So we are asking you "
            "rather than guessing.",
        ],
        "question": "Should a person be able to set a preference to always schedule the whole work "
                    "order (so the system remembers it and stops asking each time), in the first "
                    "version?",
        "options": [
            "A) YES — add the 'always schedule the whole work order' preference in the first "
            "version.",
            "B) NO — leave it out for now and keep asking each time; maybe a later version.",
            "C) Something else — please describe it.",
        ],
    },
    {
        "heading": "SCHEDULE — the schedule grid — the one-click carryover button",
        "now": [
            "The old scheduler had a one-click 'carryover' button on a scheduled job. Pressing it "
            "copied that job to the next day at the same time.",
            "The new scheduler does not have this button. The 5 August design review asked to bring "
            "it back for the first version.",
            "Your newer written description does not mention carryover at all, so we are asking you "
            "rather than guessing.",
        ],
        "question": "Should the one-click carryover button — copy a job to the next day at the same "
                    "time — be brought back in the first version?",
        "options": [
            "A) YES — bring the carryover button back in the first version.",
            "B) NO — leave it out for now.",
            "C) Something else — please describe it.",
        ],
    },
    {
        "heading": "SCHEDULE — the schedule grid — what the carryover button is called",
        "now": [
            "This one is only about what the button is CALLED, if it comes back (see the previous "
            "question).",
            "The old name was 'Carryover'. The design review suggested a clearer name — 'Add a Day' "
            "or 'Extend a Day' — but the final wording is not decided.",
        ],
        "question": "If the button comes back, what should it be called?",
        "options": [
            "A) 'Add a Day'",
            "B) 'Extend a Day'",
            "C) Keep 'Carryover', or something else — please write the exact words you want.",
        ],
    },
    {
        "heading": "SCHEDULE — the schedule grid — carryover on a job that runs several days",
        "now": [
            "Some jobs are scheduled across several days in a row.",
            "When you press carryover on one of these, there are two ways it could behave: it could "
            "add just ONE more day at the end, or it could copy the whole run of days again.",
            "The design review said it should add just one more day. Your newer written description "
            "does not cover this, so we are checking with you.",
        ],
        "question": "When you press carryover on a job that already runs across several days, should "
                    "it add just one more day, or copy the whole run of days again?",
        "options": [
            "A) Add just ONE more day at the end.",
            "B) Copy the whole run of days again.",
            "C) Something else — please describe it.",
        ],
    },
    {
        "heading": "SCHEDULE — the schedule grid — week view — dragging a job to the next day",
        "now": [
            "In week view you can see several days at once.",
            "The design review suggested that — as well as, or instead of, a carryover button — a "
            "person should be able to drag a scheduled job straight onto the next day.",
            "Your newer written description does not mention this, so it is your call.",
        ],
        "question": "In week view, should a person be able to drag a scheduled job straight onto the "
                    "next day, as another way to carry it over?",
        "options": [
            "A) YES — allow dragging a job onto the next day in week view.",
            "B) NO — keep a carryover button as the only way.",
            "C) Something else — please describe it.",
        ],
    },
]

BRANKO = {
    "path": "build/filters/questions-2026-08-17/Filters-and-Schedule_Questions-for-Branko_2026-08-17.docx",
    "title": "Questions for Branko Cicovic",
    "subtitle": "Filters and Schedule — Product Owner: Branko Cicovic",
    "intro": (
        "Hello Branko — this is everything we have open across your projects FILTERS and "
        "SCHEDULE, gathered into one place so you can answer it in one sitting instead of a "
        "trickle of separate messages. SHORT ANSWERS ARE PERFECT — a letter, or one line. "
        "Nothing here needs an essay.\n"
        "There are EIGHT questions in total. The first THREE (1 to 3) are the ones we already "
        "sent you — they are repeated here only so everything is in one place. Questions 4 to 8 "
        "are NEW FOLLOW-UPS, all about the Schedule's 'carryover' feature and one scheduling "
        "preference from the design review — so if you have already answered 1 to 3, you only "
        "need to look at 4 to 8.\n"
        "Every question says which project and screen it is about, because you look after "
        "Filters, Schedule and Global Search. Each one is a point where two of your own "
        "documents disagree, or where your written description does not yet say the thing we "
        "need — so we are asking you which to keep, rather than guessing. To be clear: we have "
        "not edited any of your tickets or descriptions.\n"
        "The first question has two tests parked on it right now, so it is the one that "
        "unblocks work."
    ),
    "questions": [
        {
            "heading": "FILTERS — the Work Orders list — the Status button on the Estimates and "
                       "Completed tabs",
            "now": [
                "Two of our tests are on hold on this one point, and two answers are on record "
                "that disagree with each other.",
                "The Work Orders list has tabs across the top. Two of them — Estimates and "
                "Completed — already show only one kind of work order. There is also a row of "
                "filter buttons, and one of them is Status.",
                "Your written description says the Status button is NOT SHOWN AT ALL on those "
                "two tabs.",
                "But you told us on 17 July that the Status button IS SHOWN, greyed out, "
                "already filled in with that tab's own status, and cannot be changed. Our QA "
                "lead agreed with that on 30 July, and the design shows it that way too.",
                "Why we are asking rather than choosing: we have set the two tests to your July "
                "answer, because that is what you and our QA lead actually decided — but the "
                "written description still says the opposite. So one of them has to change, and "
                "it is your call which.",
            ],
            "question": "On the Estimates and Completed tabs, is the Status button hidden, or "
                        "shown greyed out and already filled in?",
            "options": [
                "A) NOT SHOWN AT ALL on those two tabs — the written description is right, and "
                "my July answer is out of date.",
                "B) SHOWN, GREYED OUT AND ALREADY FILLED IN — my July answer stands, and the "
                "description needs correcting. (Then we will also raise it so the product can be "
                "fixed to match.)",
                "C) Something else — please describe it.",
            ],
        },
        {
            "heading": "FILTERS — the filter buttons on the Parts pages and the Report pages",
            "now": [
                "The redesign puts a row of filter buttons on the Parts pages and on the Report "
                "pages. Your written description says the filters those pages already had are "
                "moved into the new row — but it does not list exactly WHICH buttons should "
                "appear on WHICH page.",
                "Your engineering team was going to send us that page-by-page list, and it has "
                "not arrived yet. Without it we can check that the buttons on those pages work, "
                "but we cannot yet check that each page is showing exactly the right set of "
                "buttons — so those tests say 'confirm the exact buttons later' and are waiting "
                "on this.",
                "Why we are asking you: it is a product decision — which filter buttons belong "
                "on each page — and you are the person who can confirm it.",
            ],
            "question": "For the Parts pages and the Report pages, can you confirm which filter "
                        "buttons should appear on each page?",
            "options": [
                "A) Every page keeps exactly the same filters it had before the redesign — "
                "nothing added or removed — so the old set for each page is the answer.",
                "B) There is a specific page-by-page list — you (or engineering) will send it, "
                "and we will check each page against it.",
                "C) Something else — please describe it.",
            ],
        },
        {
            "heading": "SCHEDULE — the pop-up window that opens when you click a scheduled job",
            "now": [
                "When someone clicks a job on the schedule, a small pop-up window opens with "
                "actions on it. Right now that window offers DELETE only.",
                "An earlier version of your description mentioned a REASSIGN action in that "
                "window, and a later version took it out. Separately, a job can already be moved "
                "to a different technician by DRAGGING it on the calendar.",
                "So there are two sensible possibilities and we do not want to guess: either the "
                "window is correct with Delete only and reassigning is done by dragging, or the "
                "window should also offer a Reassign action. We have kept the test flagged and "
                "left it for your decision.",
            ],
            "question": "In that pop-up window, should there be a REASSIGN action, or is Delete "
                        "the only action and reassigning is done by dragging the job to another "
                        "technician?",
            "options": [
                "A) DELETE ONLY — reassigning is done by dragging the job to another technician. "
                "The window is correct as it is.",
                "B) ADD A REASSIGN ACTION to the window as well.",
                "C) Something else — please describe it.",
            ],
        },
    ] + FOLLOWUP_5,
}


BRANKO_FOLLOWUP = {
    "path": "build/filters/questions-2026-08-17/Filters-and-Schedule_Questions-for-Branko_FOLLOW-UP-5_2026-08-17.docx",
    "title": "Follow-up questions for Branko Cicovic",
    "subtitle": "Schedule — Product Owner: Branko Cicovic",
    "start": 4,
    "intro": (
        "Hello Branko — these are five short FOLLOW-UP questions (numbers 4 to 8) that go with the "
        "three we sent you on 17 August. They are all about the Schedule, and all about one thing: "
        "the 'carryover' feature from the design review (copying a scheduled job to the next day), "
        "plus one scheduling preference. About five minutes. SHORT ANSWERS ARE PERFECT — a letter, "
        "or one line.\n"
        "The design review listed each of these for the first version, but your newer written "
        "description does not mention them — so we are asking you rather than guessing. To be clear: "
        "we have not edited any of your tickets or descriptions."
    ),
    "questions": FOLLOWUP_5,
}


def build(doc_spec):
    start = doc_spec.get("start", 1)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title
    t = doc.add_heading(doc_spec["title"], level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle
    sub = doc.add_paragraph()
    r = sub.add_run(doc_spec["subtitle"])
    r.bold = True
    r.font.size = Pt(12)

    # Intro (may contain multiple paragraphs separated by \n)
    for para in doc_spec["intro"].split("\n"):
        doc.add_paragraph(para)

    doc.add_paragraph("")  # spacer

    for i, q in enumerate(doc_spec["questions"], start=start):
        # Numbered question heading
        h = doc.add_heading(level=1)
        hr = h.add_run(f"Question {i} — {q['heading']}")

        # What happens now
        wl = doc.add_paragraph()
        wl.add_run("What happens now").bold = True
        for line in q["now"]:
            p = doc.add_paragraph(line)
            p.paragraph_format.left_indent = Pt(18)

        # The question
        ql = doc.add_paragraph()
        ql.add_run("The question").bold = True
        qp = doc.add_paragraph(q["question"])
        qp.paragraph_format.left_indent = Pt(18)

        # Options (bulleted)
        ol = doc.add_paragraph()
        ol.add_run("Options").bold = True
        for opt in q["options"]:
            doc.add_paragraph(opt, style="List Bullet")

        # Your answer
        ap = doc.add_paragraph()
        ap.add_run("Your answer:").bold = True
        blank = doc.add_paragraph("________________________________________________________________")
        blank.paragraph_format.left_indent = Pt(18)
        # extra typing room
        doc.add_paragraph("")
        doc.add_paragraph("")

    doc.save(doc_spec["path"])
    return doc_spec["path"]


for spec in (CHRIS, BRANKO, BRANKO_FOLLOWUP):
    p = build(spec)
    print("wrote", p, "questions:", len(spec["questions"]))
